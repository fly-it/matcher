# import streamlit as st
# import os
# import PyPDF2
# import docx
# import utils.functions as f
# import math
# import json

# cwd = os.getcwd()

# # Function to extract text from PDF
# def extract_text_from_pdf(file_path):
#     with open(file_path, "rb") as f:
#         reader = PyPDF2.PdfReader(f)
#         text = [page.extract_text() for page in reader.pages]
#     return "\n".join(text)

# # Function to extract text from DOCX
# def extract_text_from_docx(file_path):
#     doc = docx.Document(file_path)
#     text = [p.text for p in doc.paragraphs]
#     return "\n".join(text)

# # Function to list files in a directory
# def list_files_in_directory():
#     directory = os.path.join(cwd,"data","resumes")
#     files = []
#     if os.path.exists(directory):
#         files = os.listdir(directory)
#         files = [file for file in files if file.endswith(('.pdf', '.docx'))]
#     return files

# # Function to save uploaded file
# def save_uploaded_file(file):
#     directory = os.path.join(cwd,"data","resumes")
#     with open(os.path.join(directory, file.name), "wb") as f:
#         f.write(file.getbuffer())
#         st.write(f)
#     return st.success(f"Saved file: {file.name} in {directory}")

# # Placeholder for the resume comparison function
# def compare_resume_to_job_description(resume_text, job_description):
#     # You will implement this function
#     pass
# # tab1, tab2, tab3 = st.tabs(["Similarité", "Recherche", "Autre"])

# # Streamlit app layout
# st.title("Fly'IT Matcher")

# # Retrieve and list existing resumes
# # existing_files = list_files_in_directory()


# # uploaded_files = st.file_uploader("Ajouter un nouveau CV", accept_multiple_files=False, type=['pdf', 'docx'])
# # # Save uploaded files
# # if uploaded_files:
# #     for uploaded_file in uploaded_files:
# #         save_uploaded_file(uploaded_file)
# #         existing_files.append(uploaded_file.name)  # Add the new file to the existing files list

# existing_resumes, resumes_names = f.get_resumes_from_firestore()        

# # st.subheader('CVs disponibles')
# resume_selected = st.selectbox('Sélectioner un CV pour la comparaison:', resumes_names)
# # resume_data = existing_resumes[resume_selected]
# # Job description text area
# job_description = st.text_area("Coller la description de  la mission ici:")

# # Button to trigger comparison
# if st.button('Comparer'):
#     if job_description and resume_selected:
#         # file_path = os.path.join(cwd,"data","resumes", resume_selected)
#         resume_data = existing_resumes[resume_selected]

#         resume_text = resume_data
#         # if resume_selected.endswith('.pdf'):
#         #     resume_text = extract_text_from_pdf(file_path)
#         # elif resume_selected.endswith('.docx'):
#         #     resume_text = extract_text_from_docx(file_path)
#         results = compare_resume_to_job_description(resume_text, job_description)
#         result_placeholder = st.subheader("Le résultat de la comparaison va aparaitre ici")

#         with st.spinner("Analyse en cours..."):
#             # result_placeholder  = st.subheader("Similarité: ")
#             job_description_enhanced = f.process_job_description(job_description)
#             resume_enhanced = f.process_resume(resume_text)
            
#             # st.header(f"{similarity}%")
#             with st.container():
#                 processed_resume_column, processed_jobDescription_column = st.columns(2)
#                 with processed_jobDescription_column:
#                     st.subheader("L'essentiel de la mission")
#                     st.write(job_description_enhanced)
#                 with processed_resume_column:
#                     st.subheader("L'essentiel du CV")
#                     st.write(resume_enhanced)
#             # st.success("Terminer")
#         resume_cleaned_text = f.clean_text(resume_enhanced)
#         job_description_cleaned_text = f.clean_text(job_description_enhanced)
        
#         resume_embedding = f.embed_with_miniLM(resume_cleaned_text)
#         job_description_embedding = f.embed_with_miniLM(job_description_cleaned_text)
#         similarity = math.floor(f.compare_embeddings(resume_embedding, job_description_embedding) * 100)
#         result_placeholder.subheader(f"Le score de similarité : {similarity}%")
#     else:
#         st.warning("Veuillez sélectionner un CV et mettre une description d'une mission pour comparer!")
######################################################################################################################################

import streamlit as st
import os
import PyPDF2
import docx
import utils.functions as f
import math

cwd = os.getcwd()

# Function to extract text from PDF
def extract_text_from_pdf(file_path):
    with open(file_path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        text = [page.extract_text() for page in reader.pages]
    return "\n".join(text)

# Function to extract text from DOCX
def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    text = [p.text for p in doc.paragraphs]
    return "\n".join(text)

# Function to list files in a directory
def list_files_in_directory():
    directory = os.path.join(cwd,"data","resumes")
    files = []
    if os.path.exists(directory):
        files = os.listdir(directory)
        files = [file for file in files if file.endswith(('.pdf', '.docx'))]
    return files

# Function to save uploaded file
def save_uploaded_file(file):
    directory = os.path.join(cwd,"data","resumes")
    with open(os.path.join(directory, file.name), "wb") as f:
        f.write(file.getbuffer())
        st.write(f)
    return st.success(f"Saved file: {file.name} in {directory}")

# Placeholder for the resume comparison function
def compare_resume_to_job_description(resume_text, job_description):
    # You will implement this function
    pass
# tab1, tab2, tab3 = st.tabs(["Similarité", "Recherche", "Autre"])

# Streamlit app layout
st.title("Fly'IT Matcher")

# Retrieve and list existing resumes
existing_files = list_files_in_directory()

# # Sidebar for new file upload
# st.sidebar.header('Ajouter un nouveau CV')
# uploaded_files = st.sidebar.file_uploader("Ajouter", accept_multiple_files=True, type=['pdf', 'docx'])

# # Save uploaded files
# if uploaded_files:
#     for uploaded_file in uploaded_files:
#         save_uploaded_file(uploaded_file)
#         existing_files.append(uploaded_file.name)  # Add the new file to the existing files list

# # Update the list of existing files
# sb = st.sidebar.header('CVs disponibles')
# resume_selected = st.sidebar.selectbox('Sélectioner un CV pour la comparaison:', existing_files)
# st.subheader("Ajouter un nouveau CV")
uploaded_files = st.file_uploader("Ajouter un nouveau CV", accept_multiple_files=False, type=['pdf', 'docx'])
# Save uploaded files
if uploaded_files:
    for uploaded_file in uploaded_files:
        save_uploaded_file(uploaded_file)
        existing_files.append(uploaded_file.name)  # Add the new file to the existing files list

# st.subheader('CVs disponibles')
resume_selected = st.selectbox('Sélectioner un CV pour la comparaison:', existing_files)
# Job description text area
job_description = st.text_area("Coller la description de  la mission ici:")

# Button to trigger comparison
if st.button('Comparer'):
    if job_description and resume_selected:
        file_path = os.path.join(cwd,"data","resumes", resume_selected)
        resume_text = ""
        if resume_selected.endswith('.pdf'):
            resume_text = extract_text_from_pdf(file_path)
        elif resume_selected.endswith('.docx'):
            resume_text = extract_text_from_docx(file_path)
        results = compare_resume_to_job_description(resume_text, job_description)
        result_placeholder = st.subheader("Le résultat de la comparaison va aparaitre ici")

        with st.spinner("Analyse en cours..."):
            # result_placeholder  = st.subheader("Similarité: ")
            job_description_enhanced = f.process_job_description(job_description)
            resume_enhanced = f.process_resume(resume_text)
            
            # st.header(f"{similarity}%")
            with st.container():
                processed_resume_column, processed_jobDescription_column = st.columns(2)
                with processed_jobDescription_column:
                    st.subheader("L'essentiel de la mission")
                    st.write(job_description_enhanced)
                with processed_resume_column:
                    st.subheader("L'essentiel du CV")
                    st.write(resume_enhanced)
            # st.success("Terminer")
        resume_cleaned_text = f.clean_text(resume_enhanced)
        job_description_cleaned_text = f.clean_text(job_description_enhanced)
        
        resume_embedding = f.embed_with_miniLM(resume_cleaned_text)
        job_description_embedding = f.embed_with_miniLM(job_description_cleaned_text)
        similarity = math.floor(f.compare_embeddings(resume_embedding, job_description_embedding) * 100)
        result_placeholder.subheader(f"Le score de similarité : {similarity}%")
    else:
        st.warning("Veuillez sélectionner un CV et mettre une description d'une mission pour comparer!")
# with tab2:
#     # Streamlit app layout
#     st.title("Chercher les offres d'emplois")

#     # Retrieve and list existing resumes
#     existing_files = list_files_in_directory()

#     # Sidebar for new file upload
#     sb = st.sidebar.header('Chercher les offres')